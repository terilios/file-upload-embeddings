from typing import Dict, Any, List
import re
from datetime import datetime
import json
from pathlib import Path
import spacy
import magic
from langdetect import detect
import PyPDF2
import docx
import email
from email import policy
from email.parser import BytesParser

# Load spaCy model for entity recognition
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    import subprocess
    subprocess.run(["python", "-m", "spacy", "download", "en_core_web_sm"])
    nlp = spacy.load("en_core_web_sm")

async def extract_metadata(
    content: bytes,
    filename: str,
    content_type: str
) -> Dict[str, Any]:
    """
    Extract metadata from document content.
    
    Args:
        content: Raw document content
        filename: Original filename
        content_type: MIME type of the document
    
    Returns:
        Dictionary containing extracted metadata
    """
    metadata = {
        "filename": filename,
        "content_type": content_type,
        "file_size": len(content),
        "processed_at": datetime.utcnow().isoformat(),
    }
    
    # Detect file type using magic
    file_type = magic.from_buffer(content, mime=True)
    metadata["detected_type"] = file_type
    
    try:
        # Extract text based on file type
        if file_type == "application/pdf":
            text, pdf_metadata = extract_pdf_metadata(content)
            metadata.update(pdf_metadata)
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text, docx_metadata = extract_docx_metadata(content)
            metadata.update(docx_metadata)
        elif file_type == "message/rfc822":
            text, email_metadata = extract_email_metadata(content)
            metadata.update(email_metadata)
        else:
            text = content.decode('utf-8', errors='ignore')
        
        # Add text analysis metadata
        metadata.update(analyze_text_content(text))
        
    except Exception as e:
        metadata["extraction_error"] = str(e)
        text = content.decode('utf-8', errors='ignore')
    
    return metadata

def extract_pdf_metadata(content: bytes) -> tuple[str, Dict[str, Any]]:
    """Extract text and metadata from PDF content."""
    metadata = {}
    text = ""
    
    try:
        with BytesIO(content) as stream:
            reader = PyPDF2.PdfReader(stream)
            
            # Extract text
            text = "\n".join(
                page.extract_text()
                for page in reader.pages
            )
            
            # Get PDF metadata
            info = reader.metadata
            if info:
                metadata["pdf_metadata"] = {
                    "title": info.get("/Title", ""),
                    "author": info.get("/Author", ""),
                    "subject": info.get("/Subject", ""),
                    "creator": info.get("/Creator", ""),
                    "producer": info.get("/Producer", ""),
                    "creation_date": info.get("/CreationDate", ""),
                    "modification_date": info.get("/ModDate", "")
                }
            
            metadata["page_count"] = len(reader.pages)
            
    except Exception as e:
        metadata["pdf_extraction_error"] = str(e)
    
    return text, metadata

def extract_docx_metadata(content: bytes) -> tuple[str, Dict[str, Any]]:
    """Extract text and metadata from DOCX content."""
    metadata = {}
    text = ""
    
    try:
        with BytesIO(content) as stream:
            doc = docx.Document(stream)
            
            # Extract text
            text = "\n".join(
                paragraph.text
                for paragraph in doc.paragraphs
            )
            
            # Get core properties
            core_props = doc.core_properties
            metadata["docx_metadata"] = {
                "author": core_props.author or "",
                "title": core_props.title or "",
                "subject": core_props.subject or "",
                "created": core_props.created.isoformat() if core_props.created else "",
                "modified": core_props.modified.isoformat() if core_props.modified else "",
                "last_modified_by": core_props.last_modified_by or ""
            }
            
            # Count sections, tables, etc.
            metadata["paragraph_count"] = len(doc.paragraphs)
            metadata["section_count"] = len(doc.sections)
            metadata["table_count"] = len(doc.tables)
            
    except Exception as e:
        metadata["docx_extraction_error"] = str(e)
    
    return text, metadata

def extract_email_metadata(content: bytes) -> tuple[str, Dict[str, Any]]:
    """Extract text and metadata from email content."""
    metadata = {}
    
    try:
        msg = BytesParser(policy=policy.default).parsebytes(content)
        
        # Extract email metadata
        metadata["email_metadata"] = {
            "subject": msg.get("subject", ""),
            "from": msg.get("from", ""),
            "to": msg.get("to", ""),
            "cc": msg.get("cc", ""),
            "date": msg.get("date", "")
        }
        
        # Get email body
        if msg.is_multipart():
            text = ""
            for part in msg.walk():
                if part.get_content_type() == "text/plain":
                    text += part.get_content()
        else:
            text = msg.get_content()
            
    except Exception as e:
        metadata["email_extraction_error"] = str(e)
        text = content.decode('utf-8', errors='ignore')
    
    return text, metadata

def analyze_text_content(text: str) -> Dict[str, Any]:
    """Analyze text content for additional metadata."""
    metadata = {}
    
    try:
        # Detect language
        metadata["language"] = detect(text)
        
        # Basic text statistics
        metadata["character_count"] = len(text)
        metadata["word_count"] = len(text.split())
        metadata["line_count"] = len(text.splitlines())
        
        # Process with spaCy for entity recognition
        doc = nlp(text[:1000000])  # Limit text length for processing
        
        # Extract named entities
        entities = {}
        for ent in doc.ents:
            if ent.label_ not in entities:
                entities[ent.label_] = []
            if ent.text not in entities[ent.label_]:
                entities[ent.label_].append(ent.text)
        
        metadata["named_entities"] = entities
        
        # Detect document structure
        structure = detect_document_structure(text)
        metadata["document_structure"] = structure
        
    except Exception as e:
        metadata["analysis_error"] = str(e)
    
    return metadata

def detect_document_structure(text: str) -> Dict[str, Any]:
    """Detect document structure patterns."""
    structure = {
        "has_headers": False,
        "has_lists": False,
        "has_tables": False,
        "has_code_blocks": False,
        "section_count": 0
    }
    
    # Header detection patterns
    header_patterns = [
        r'^#{1,6}\s+.+$',  # Markdown headers
        r'^[A-Z][A-Za-z\s]+:$',  # Capitalized labels
        r'^\d+\.\s+[A-Z][A-Za-z\s]+$',  # Numbered sections
        r'^[A-Z][A-Z\s]+$'  # All caps headers
    ]
    
    # List detection patterns
    list_patterns = [
        r'^\s*[-*+]\s+.+$',  # Unordered lists
        r'^\s*\d+\.\s+.+$'   # Ordered lists
    ]
    
    # Table detection patterns
    table_patterns = [
        r'^\|.+\|$',  # Markdown tables
        r'^\s*[|+].+[|+]\s*$'  # ASCII tables
    ]
    
    # Code block detection patterns
    code_patterns = [
        r'```[\s\S]*?```',  # Markdown code blocks
        r'    [\w\s]+$'     # Indented code blocks
    ]
    
    lines = text.split('\n')
    
    for line in lines:
        # Check headers
        if any(re.match(pattern, line) for pattern in header_patterns):
            structure["has_headers"] = True
            structure["section_count"] += 1
        
        # Check lists
        if any(re.match(pattern, line) for pattern in list_patterns):
            structure["has_lists"] = True
        
        # Check tables
        if any(re.match(pattern, line) for pattern in table_patterns):
            structure["has_tables"] = True
        
    # Check code blocks in full text
    if any(re.search(pattern, text, re.MULTILINE) for pattern in code_patterns):
        structure["has_code_blocks"] = True
    
    return structure
